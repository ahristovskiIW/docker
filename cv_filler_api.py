"""
FastAPI CV Filler - Fills an empty DOCX CV template with data from JSON
"""

import json
import os
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel

app = FastAPI(title="CV Filler API", version="1.0.0")


class PersonalInfo(BaseModel):
    name: str
    email: Optional[str] = ""
    phone: Optional[str] = ""
    address: Optional[str] = ""
    linkedin: Optional[str] = ""
    github: Optional[str] = ""
    website: Optional[str] = ""
    additional: Optional[Dict[str, Any]] = {}


class Education(BaseModel):
    institution: str
    degree: str
    field_of_study: str
    start_date: Optional[str] = ""
    end_date: Optional[str] = ""
    description: Optional[str] = ""


class Project(BaseModel):
    name: str
    description: str
    technologies: List[str]
    start_date: Optional[str] = ""
    end_date: Optional[str] = ""
    additional: Optional[Dict[str, Any]] = {}


class Language(BaseModel):
    language: str
    proficiency: Optional[str] = ""


class CVData(BaseModel):
    personal_info: PersonalInfo
    education: List[Education]
    certificates: List[str] = []
    soft_skills: List[str]
    programming_skills: List[str]
    projects: List[Project]
    languages: List[Union[str, Language]] = []
    other_info: str


class CVFiller:
    """Handles filling the DOCX CV template with JSON data"""

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = Document(template_path)

    def fill_cv(self, cv_data: CVData) -> Document:
        """Fill the CV template with provided data"""

        # The CV is structured as a single table with specific cells
        if not self.doc.tables:
            raise ValueError("Template document does not contain any tables")

        table = self.doc.tables[0]

        # Fill header (Row 0, Cell 1) - Name and Position
        self._fill_header(table, cv_data.personal_info)

        # Fill Summary (Row 1, Cells 0 and 1)
        self._fill_summary(table, cv_data.other_info)

        # Fill Education and Industry Knowledge (Row 1, Cell 2)
        self._fill_education_and_industry(table, cv_data.education, cv_data.projects)

        # Fill Skills (Row 2, Cell 2)
        self._fill_skills(table, cv_data.programming_skills, cv_data.soft_skills)

        # Fill Professional Experience (Row 2-3, Cells 0-1)
        self._fill_professional_experience(table, cv_data.projects)

        return self.doc

    def _fill_header(self, table, personal_info: PersonalInfo):
        """Fill the header with name and position"""
        cell = table.rows[0].cells[1]
        cell.text = ""  # Clear existing content

        # Add name
        p = cell.add_paragraph()
        run = p.add_run(personal_info.name)
        run.bold = True
        run.font.size = Pt(16)

        # Determine position from additional info or use default
        position = personal_info.additional.get("position", "Integration Developer")
        p = cell.add_paragraph()
        run = p.add_run(position)
        run.font.size = Pt(11)

    def _fill_summary(self, table, summary_text: str):
        """Fill the summary section"""
        # Summary appears in both cell[1,0] and cell[1,1] (merged cells)
        for cell_idx in [0, 1]:
            cell = table.rows[1].cells[cell_idx]
            cell.text = ""

            # Add SUMMARY header
            p = cell.add_paragraph()
            run = p.add_run("SUMMARY")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(192, 0, 0)

            # Add summary text
            p = cell.add_paragraph()
            run = p.add_run(summary_text)
            run.font.size = Pt(9)

    def _fill_education_and_industry(
        self, table, education: List[Education], projects: List[Project]
    ):
        """Fill education and industry knowledge section"""
        cell = table.rows[1].cells[2]
        cell.text = ""

        # EDUCATION header
        p = cell.add_paragraph()
        run = p.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(192, 0, 0)

        # Add education entries
        for edu in education:
            p = cell.add_paragraph()
            edu_text = f"{edu.degree} in {edu.field_of_study} - {edu.institution}"
            run = p.add_run(edu_text)
            run.font.size = Pt(9)

        # Add spacing
        cell.add_paragraph()
        cell.add_paragraph()

        # INDUSTRY KNOWLEDGE header
        p = cell.add_paragraph()
        run = p.add_run("INDUSTRY KNOWLEDGE")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(192, 0, 0)

        # Extract unique industries from projects
        industries = set()
        for project in projects:
            if "industry" in project.additional:
                industries.add(project.additional["industry"])

        # Add industries
        for industry in sorted(industries):
            p = cell.add_paragraph()
            run = p.add_run(industry)
            run.font.size = Pt(9)

    def _fill_skills(
        self, table, programming_skills: List[str], soft_skills: List[str]
    ):
        """Fill the skills section"""
        cell = table.rows[2].cells[2]
        cell.text = ""

        # SKILLS header
        p = cell.add_paragraph()
        run = p.add_run("SKILLS")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(192, 0, 0)

        # Group programming skills by category
        skill_groups = self._categorize_skills(programming_skills)

        # Add skill categories
        for category, skills in skill_groups.items():
            p = cell.add_paragraph()
            skill_text = f"{category}: {', '.join(skills)}"
            run = p.add_run(skill_text)
            run.font.size = Pt(9)

        # Add soft skills
        if soft_skills:
            p = cell.add_paragraph()
            run = p.add_run(f"Soft Skills: {', '.join(soft_skills[:3])}")
            run.font.size = Pt(9)

    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        """Categorize programming skills into logical groups"""
        categories = {
            "Programming Language": [],
            "Backend Development": [],
            "Frontend Development": [],
            "Cloud & DevOps": [],
            "Database": [],
            "Web Services": [],
            "CI/CD": [],
            "Supporting tools": [],
            "Methodologies": [],
        }

        # Define categorization rules
        language_keywords = ["C#", "Java", "JavaScript", "TypeScript", "Python"]
        backend_keywords = [".NET Core", "Spring Boot", "Node.js"]
        frontend_keywords = ["Angular", "React", "Vue", "HTML", "CSS", "Bootstrap"]
        cloud_keywords = ["AWS", "Azure", "Docker", "Kubernetes"]
        db_keywords = ["PostgreSQL", "MySQL", "MSSQL", "MongoDB", "Oracle"]
        webservice_keywords = ["REST", "SOAP", "GraphQL"]
        cicd_keywords = ["Jenkins", "Azure DevOps", "CI/CD", "GitHub Actions"]
        tools_keywords = ["Git", "Jira", "Maven", "NPM", "SonarQube", "Swagger"]
        integration_keywords = ["SnapLogic", "MuleSoft", "Boomi"]

        for skill in skills:
            # Check each category
            if any(kw.lower() in skill.lower() for kw in language_keywords):
                categories["Programming Language"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in backend_keywords):
                categories["Backend Development"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in frontend_keywords):
                categories["Frontend Development"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in cloud_keywords):
                categories["Cloud & DevOps"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in db_keywords):
                categories["Database"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in webservice_keywords):
                categories["Web Services"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in cicd_keywords):
                categories["CI/CD"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in tools_keywords):
                categories["Supporting tools"].append(skill)
            elif any(kw.lower() in skill.lower() for kw in integration_keywords):
                categories["Backend Development"].append(skill)
            else:
                categories["Supporting tools"].append(skill)

        # Remove empty categories
        return {k: v for k, v in categories.items() if v}

    def _fill_professional_experience(self, table, projects: List[Project]):
        """Fill the professional experience section"""
        # Professional experience spans rows 2-3, cells 0-1
        # We'll use row 2, cells 0-1 for the content
        for cell_idx in [0, 1]:
            cell = table.rows[2].cells[cell_idx]
            cell.text = ""

            # Add PROFESSIONAL EXPERIENCE header
            p = cell.add_paragraph()
            run = p.add_run("PROFESSIONAL EXPERIENCE")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(192, 0, 0)

            # Add projects in reverse chronological order (most recent first)
            sorted_projects = sorted(
                projects, key=lambda x: self._parse_date(x.start_date), reverse=True
            )

            for project in sorted_projects:
                # Add spacing
                cell.add_paragraph()

                # Project name and dates
                p = cell.add_paragraph()
                # Only include date range if there's actual date information
                if project.start_date or project.end_date:
                    date_range = f"{project.start_date} - {project.end_date}"
                    run = p.add_run(f"{project.name} ({date_range})")
                else:
                    run = p.add_run(project.name)
                run.bold = True
                run.font.size = Pt(9)

                # Project description
                p = cell.add_paragraph()
                run = p.add_run(project.description)
                run.font.size = Pt(9)

                # Technologies
                if project.technologies:
                    p = cell.add_paragraph()
                    tech_text = f"Technologies: {', '.join(project.technologies)}"
                    run = p.add_run(tech_text)
                    run.font.size = Pt(8)
                    run.italic = True

    def _parse_date(self, date_str: str) -> datetime:
        """Parse date string for sorting"""
        if not date_str or date_str.lower() == "current":
            return datetime.now()
        try:
            return datetime.strptime(date_str, "%B %Y")
        except:
            return datetime.now()

    def save(self, output_path: str):
        """Save the filled document"""
        self.doc.save(output_path)
        return output_path


# API Endpoints


@app.post("/fill-cv/")
async def fill_cv_from_json(
    template: UploadFile = File(..., description="Empty DOCX CV template"),
    cv_data_json: UploadFile = File(..., description="JSON file with CV data"),
):
    """
    Fill a CV template with data from JSON file

    - **template**: Empty DOCX CV template file
    - **cv_data_json**: JSON file containing CV data

    Returns the filled CV as a downloadable DOCX file
    """
    try:
        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Save uploaded template
            template_path = os.path.join(tmp_dir, "template.docx")
            with open(template_path, "wb") as f:
                f.write(await template.read())

            # Read and parse JSON data
            json_content = await cv_data_json.read()
            cv_data_dict = json.loads(json_content)
            cv_data = CVData(**cv_data_dict)

            # Fill the CV
            filler = CVFiller(template_path)
            filler.fill_cv(cv_data)

            # Save filled CV
            output_path = os.path.join(tmp_dir, "filled_cv.docx")
            filler.save(output_path)

            # Copy to a permanent location for download
            permanent_output = os.path.join(
                tempfile.gettempdir(),
                f"filled_cv_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            )
            filler.save(permanent_output)

            return FileResponse(
                permanent_output,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="filled_cv.docx",
            )

    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON format: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing CV: {str(e)}")


@app.post("/fill-cv-from-data/")
async def fill_cv_from_data(
    template: UploadFile = File(..., description="Empty DOCX CV template"),
    cv_data: str = Form(..., description="CV data as JSON string"),
):
    """
    Fill a CV template with data from JSON body

    - **template**: Empty DOCX CV template file
    - **cv_data**: CV data as JSON string in form field

    Returns the filled CV as a downloadable DOCX file
    """
    try:
        # Create temporary directory for processing
        with tempfile.TemporaryDirectory() as tmp_dir:
            # Save uploaded template
            template_path = os.path.join(tmp_dir, "template.docx")
            with open(template_path, "wb") as f:
                f.write(await template.read())

            # Parse JSON string to CVData object
            cv_data_dict = json.loads(cv_data)
            cv_data_obj = CVData(**cv_data_dict)

            # Fill the CV
            filler = CVFiller(template_path)
            filler.fill_cv(cv_data_obj)

            # Save filled CV
            permanent_output = os.path.join(
                tempfile.gettempdir(),
                f"filled_cv_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            )
            filler.save(permanent_output)

            return FileResponse(
                permanent_output,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename="filled_cv.docx",
            )

    except json.JSONDecodeError as e:
        raise HTTPException(status_code=400, detail=f"Invalid JSON format: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing CV: {str(e)}")


@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "CV Filler API",
        "version": "1.0.0",
        "endpoints": {
            "/fill-cv/": "Fill CV from uploaded template and JSON file",
            "/fill-cv-from-data/": "Fill CV from uploaded template and JSON body",
            "/docs": "Interactive API documentation",
        },
    }


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
