"""
Standalone CV Filler Script
Usage: python fill_cv_standalone.py
"""
from docx import Document
from docx.shared import Pt
from typing import List, Dict, Any
import json
from datetime import datetime


class SimpleCVFiller:
    """Simple CV filler without FastAPI dependencies"""

    def __init__(self, template_path: str):
        self.template_path = template_path
        self.doc = Document(template_path)

    def fill_cv(self, cv_data: dict) -> Document:
        """Fill the CV template with provided data"""

        if not self.doc.tables:
            raise ValueError("Template document does not contain any tables")

        table = self.doc.tables[0]

        # Fill different sections
        self._fill_header(table, cv_data['personal_info'])
        self._fill_summary(table, cv_data['other_info'])
        self._fill_education_and_industry(table, cv_data['education'], cv_data['projects'])
        self._fill_skills(table, cv_data['programming_skills'], cv_data['soft_skills'])
        self._fill_professional_experience(table, cv_data['projects'])

        return self.doc

    def _fill_header(self, table, personal_info: dict):
        """Fill the header with name and position"""
        cell = table.rows[0].cells[1]
        cell.text = ""

        # Add name
        p = cell.add_paragraph()
        run = p.add_run(personal_info['name'])
        run.bold = True
        run.font.size = Pt(16)

        # Add position
        position = personal_info.get('additional', {}).get('position', 'Integration Developer')
        p = cell.add_paragraph()
        run = p.add_run(position)
        run.font.size = Pt(11)

    def _fill_summary(self, table, summary_text: str):
        """Fill the summary section"""
        for cell_idx in [0, 1]:
            cell = table.rows[1].cells[cell_idx]
            cell.text = ""

            p = cell.add_paragraph()
            run = p.add_run("SUMMARY")
            run.bold = True
            run.font.size = Pt(10)

            p = cell.add_paragraph()
            run = p.add_run(summary_text)
            run.font.size = Pt(9)

    def _fill_education_and_industry(self, table, education: list, projects: list):
        """Fill education and industry knowledge section"""
        cell = table.rows[1].cells[2]
        cell.text = ""

        # Education
        p = cell.add_paragraph()
        run = p.add_run("EDUCATION")
        run.bold = True
        run.font.size = Pt(10)

        for edu in education:
            p = cell.add_paragraph()
            edu_text = f"{edu['degree']} in {edu['field_of_study']} - {edu['institution']}"
            run = p.add_run(edu_text)
            run.font.size = Pt(9)

        cell.add_paragraph()
        cell.add_paragraph()

        # Industry Knowledge
        p = cell.add_paragraph()
        run = p.add_run("INDUSTRY KNOWLEDGE")
        run.bold = True
        run.font.size = Pt(10)

        industries = set()
        for project in projects:
            if 'industry' in project.get('additional', {}):
                industries.add(project['additional']['industry'])

        for industry in sorted(industries):
            p = cell.add_paragraph()
            run = p.add_run(industry)
            run.font.size = Pt(9)

    def _fill_skills(self, table, programming_skills: list, soft_skills: list):
        """Fill the skills section"""
        cell = table.rows[2].cells[2]
        cell.text = ""

        p = cell.add_paragraph()
        run = p.add_run("SKILLS")
        run.bold = True
        run.font.size = Pt(10)

        # Categorize and add skills
        skill_groups = self._categorize_skills(programming_skills)

        for category, skills in skill_groups.items():
            p = cell.add_paragraph()
            skill_text = f"{category}: {', '.join(skills)}"
            run = p.add_run(skill_text)
            run.font.size = Pt(9)

        if soft_skills:
            p = cell.add_paragraph()
            run = p.add_run(f"Soft Skills: {', '.join(soft_skills[:3])}")
            run.font.size = Pt(9)

    def _categorize_skills(self, skills: list) -> dict:
        """Categorize programming skills"""
        categories = {
            "Programming Language": [],
            "Backend Development": [],
            "Frontend Development": [],
            "Cloud & DevOps": [],
            "Database": [],
            "Web Services": [],
            "CI/CD": [],
            "Supporting tools": []
        }

        # Categorization rules
        mappings = {
            "Programming Language": ['C#', 'Java', 'JavaScript', 'TypeScript', 'Python'],
            "Backend Development": ['.NET Core', 'Spring Boot', 'Node.js', 'SnapLogic', 'MuleSoft', 'Boomi'],
            "Frontend Development": ['Angular', 'React', 'Vue', 'HTML', 'CSS', 'Bootstrap'],
            "Cloud & DevOps": ['AWS', 'Azure', 'Docker', 'Kubernetes'],
            "Database": ['PostgreSQL', 'MySQL', 'MSSQL', 'MongoDB', 'Oracle'],
            "Web Services": ['REST', 'SOAP', 'GraphQL'],
            "CI/CD": ['Jenkins', 'Azure DevOps', 'CI/CD', 'GitHub Actions'],
            "Supporting tools": ['Git', 'Jira', 'Maven', 'NPM', 'SonarQube', 'Swagger']
        }

        for skill in skills:
            categorized = False
            for category, keywords in mappings.items():
                if any(kw.lower() in skill.lower() for kw in keywords):
                    categories[category].append(skill)
                    categorized = True
                    break
            if not categorized:
                categories["Supporting tools"].append(skill)

        return {k: v for k, v in categories.items() if v}

    def _fill_professional_experience(self, table, projects: list):
        """Fill the professional experience section"""
        for cell_idx in [0, 1]:
            cell = table.rows[2].cells[cell_idx]
            cell.text = ""

            p = cell.add_paragraph()
            run = p.add_run("PROFESSIONAL EXPERIENCE")
            run.bold = True
            run.font.size = Pt(10)

            # Sort projects by date
            sorted_projects = sorted(projects, 
                                   key=lambda x: self._parse_date(x.get('start_date', '')), 
                                   reverse=True)

            for project in sorted_projects:
                cell.add_paragraph()

                # Project header
                p = cell.add_paragraph()
                date_range = f"{project.get('start_date', '')} - {project.get('end_date', '')}"
                run = p.add_run(f"{project['name']} ({date_range})")
                run.bold = True
                run.font.size = Pt(9)

                # Description
                p = cell.add_paragraph()
                run = p.add_run(project['description'])
                run.font.size = Pt(9)

                # Technologies
                if project.get('technologies'):
                    p = cell.add_paragraph()
                    tech_text = f"Technologies: {', '.join(project['technologies'])}"
                    run = p.add_run(tech_text)
                    run.font.size = Pt(8)
                    run.italic = True

    def _parse_date(self, date_str: str) -> datetime:
        """Parse date string for sorting"""
        if not date_str or date_str.lower() == 'current':
            return datetime.now()
        try:
            return datetime.strptime(date_str, "%B %Y")
        except:
            return datetime.now()

    def save(self, output_path: str):
        """Save the filled document"""
        self.doc.save(output_path)
        return output_path


def main():
    """Main function to demonstrate usage"""

    # File paths
    template_path = "/mnt/user-data/uploads/empty_DOCX.docx"
    json_path = "/mnt/user-data/uploads/JSON_input.json"
    output_path = "/mnt/user-data/outputs/filled_cv.docx"

    # Load JSON data
    print("Loading CV data from JSON...")
    with open(json_path, 'r') as f:
        cv_data = json.load(f)

    # Fill CV
    print("Filling CV template...")
    filler = SimpleCVFiller(template_path)
    filler.fill_cv(cv_data)

    # Save output
    print(f"Saving filled CV to {output_path}...")
    filler.save(output_path)

    print("✅ CV filled successfully!")
    print(f"Output saved to: {output_path}")


if __name__ == "__main__":
    main()
 